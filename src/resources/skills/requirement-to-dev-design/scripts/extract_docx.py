# -*- coding: utf-8 -*-
"""从 Word(.docx) 提取正文与表格到文本文件，便于需求解析。依赖: pip install python-docx"""
import os
import sys

try:
    from docx import Document
except ImportError:
    print("请先安装: pip install python-docx")
    raise

def extract_docx(docx_path: str, out_path: str = None) -> str:
    docx_path = os.path.abspath(docx_path)
    if not os.path.isfile(docx_path):
        raise FileNotFoundError("文件不存在: " + docx_path)
    if out_path is None:
        base, _ = os.path.splitext(docx_path)
        out_path = base + "_原文.txt"
    doc = Document(docx_path)
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append("\t".join(cells))
        lines.append("")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return out_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python extract_docx.py <需求文档.docx> [输出路径.txt]")
        sys.exit(1)
    path = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    result = extract_docx(path, out)
    print("已写出:", result)
